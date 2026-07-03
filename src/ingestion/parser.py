import datetime
import os
import re
import uuid
import unicodedata
from collections import Counter
from typing import List, Tuple, Optional
import statistics
import logging

import fitz  # PyMuPDF
# pyrefly: ignore [missing-import]
import docx  # python-docx
# pyrefly: ignore [missing-import]
from docx.document import Document as DocxDocument
# pyrefly: ignore [missing-import]
from docx.text.paragraph import Paragraph
# pyrefly: ignore [missing-import]
from docx.table import Table

from src.ingestion.models import ParsedDocument, ParsedPage, ParsedSection, ParsedTable, TextBlock


class DocumentParsingError(Exception):
    """Base exception for parsing errors."""
    pass

class DocumentProtectedError(DocumentParsingError):
    """Raised when the document is password-protected."""
    pass

class ScannedDocumentError(DocumentParsingError):
    """Raised when the PDF has no text and only images."""
    pass

class EmptyDocumentError(DocumentParsingError):
    """Raised when the parsed document has no text content."""
    pass

class CorruptedDocumentError(DocumentParsingError):
    """Raised when the document file is corrupted or unreadable."""
    pass


class TextNormalizer:
    @staticmethod
    def normalize(text: str) -> str:
        if not text:
            return ""
        # Unicode NFKC normalization
        text = unicodedata.normalize("NFKC", text)
        # Normalize whitespace (replace tabs/multiple spaces with a single space)
        text = re.sub(r'[ \t]+', ' ', text)
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # Debounce multiple consecutive newlines (max 2)
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Bullet character normalization
        text = re.sub(r'^[•\-\*\–\u2022\u2013\u25aa\u25cf]\s*', '• ', text, flags=re.MULTILINE)
        # Fix hyphenation split across lines
        text = re.sub(r'(\w+)-\n(\w+)', r'\1\2', text)
        return text.strip()


logger = logging.getLogger(__name__)

_WATERMARK_FOOTER_RE = re.compile(
    r'Downloaded\s+by\s*:.*?on\s+\d{2}/\d{2}/\d{4}.*', re.IGNORECASE
)

def _strip_watermark(text: str) -> str:
    """Remove repeated-token watermarks and portal download footers."""
    if not text:
        return text
    # Remove "Downloaded by: ..." footer lines
    text = _WATERMARK_FOOTER_RE.sub('', text)
    # Remove tokens that appear suspiciously often on the page (watermark grid)
    words = text.split()
    if not words:
        return text
    freq = Counter(words)
    # Any token appearing more than 5 times AND making up >20% of the word count
    # is almost certainly a watermark (e.g. roll-number repeated across the page)
    total = len(words)
    noise = {tok for tok, cnt in freq.items() if cnt > 5 and cnt / total > 0.20}
    if noise:
        words = [w for w in words if w not in noise]
        text = ' '.join(words)
    return text.strip()


def _ocr_page(page) -> str:
    """Run Tesseract OCR on a single PyMuPDF page via the built-in bridge."""
    try:
        # get_textpage_ocr() calls Tesseract internally — no extra Python deps needed
        tp = page.get_textpage_ocr(language="eng", dpi=200, full=False)
        return page.get_text(textpage=tp).strip()
    except Exception as e:
        logger.warning(f"OCR failed on page: {e}")
        return ""


class PDFParser:
    def parse(self, file_path: str, doc_id: str) -> ParsedDocument:
        try:
            doc = fitz.open(file_path)
        except Exception as e:
            raise CorruptedDocumentError(f"Failed to open PDF file: {e}")

        try:
            if doc.is_encrypted:
                raise DocumentProtectedError("PDF is password-protected.")

            # Calculate font statistics for heading detection
            font_sizes = []
            for page in doc:
                dict_text = page.get_text("dict")
                for block in dict_text.get("blocks", []):
                    if block.get("type") == 0:  # text block
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                if span.get("text", "").strip():
                                    font_sizes.append(span.get("size"))

            median_font_size = statistics.median(font_sizes) if font_sizes else 10.0

            parsed_pages = []
            parsed_sections = []
            parsed_tables = []
            full_text_parts = []
            table_id_counter = 1
            section_id_counter = 1

            current_section_title = None
            current_section_level = 0
            current_section_content = []
            current_section_start_page = 1

            for page_idx, page in enumerate(doc, start=1):
                page_text_blocks = []
                page_text_parts = []
                tables_in_page = []

                # --- Decide extraction strategy for this page ---
                # Quick check: does the native text layer have real content?
                raw_native = page.get_text("text")
                clean_native = _strip_watermark(raw_native)
                has_real_text = len(clean_native.strip()) > 40

                if has_real_text:
                    # ---- Path A: native text layer ----

                    # 1. Extract tables
                    for tbl in page.find_tables():
                        extracted_tbl = tbl.extract()
                        if not extracted_tbl:
                            continue
                        headers = [str(h or "").strip() for h in extracted_tbl[0]]
                        rows = [[str(cell or "").strip() for cell in row] for row in extracted_tbl[1:]]
                        raw_text = "\n".join([" | ".join(row) for row in extracted_tbl])
                        p_table = ParsedTable(
                            table_id=f"T{table_id_counter}",
                            headers=headers,
                            rows=rows,
                            source_page=page_idx,
                            raw_text=raw_text
                        )
                        parsed_tables.append(p_table)
                        tables_in_page.append(p_table)
                        table_id_counter += 1

                    # 2. Extract text blocks with heading detection
                    dict_text = page.get_text("dict")
                    blocks_with_coords = []
                    for b in dict_text.get("blocks", []):
                        if b.get("type") == 0:
                            bbox = b.get("bbox")
                            blocks_with_coords.append((bbox[1], bbox[0], b))
                    blocks_with_coords.sort(key=lambda item: (round(item[0] / 5) * 5, item[1]))

                    for b in [item[2] for item in blocks_with_coords]:
                        block_text_parts = []
                        block_font_sizes = []
                        block_is_bold = False
                        for line in b.get("lines", []):
                            for span in line.get("spans", []):
                                span_text = span.get("text", "")
                                if span_text.strip():
                                    block_text_parts.append(span_text)
                                    block_font_sizes.append(span.get("size"))
                                    if bool(span.get("flags") & 2):
                                        block_is_bold = True
                        if not block_text_parts:
                            continue

                        block_text = _strip_watermark(" ".join(block_text_parts))
                        if not block_text:
                            continue

                        avg_font_size = sum(block_font_sizes) / len(block_font_sizes) if block_font_sizes else median_font_size
                        is_heading = False
                        heading_level = 0
                        if avg_font_size > median_font_size * 1.15:
                            is_heading = True
                            if avg_font_size > median_font_size * 1.5:
                                heading_level = 1
                            elif avg_font_size > median_font_size * 1.3:
                                heading_level = 2
                            else:
                                heading_level = 3
                        elif avg_font_size > median_font_size * 1.05 and block_is_bold:
                            is_heading = True
                            heading_level = 3

                        block_type = "heading" if is_heading else "text"
                        if not is_heading and block_text.strip().startswith(('•', '-', '*')):
                            block_type = "bullet"

                        t_block = TextBlock(
                            text=block_text,
                            font_size=avg_font_size,
                            is_bold=block_is_bold,
                            bbox=b.get("bbox"),
                            block_type=block_type
                        )
                        page_text_blocks.append(t_block)
                        page_text_parts.append(block_text)

                        if is_heading:
                            if current_section_content or current_section_title:
                                parsed_sections.append(ParsedSection(
                                    section_id=f"S{section_id_counter}",
                                    heading_text=current_section_title,
                                    heading_level=current_section_level,
                                    content=TextNormalizer.normalize("\n".join(current_section_content)),
                                    start_page=current_section_start_page
                                ))
                                section_id_counter += 1
                            current_section_title = block_text.strip()
                            current_section_level = heading_level
                            current_section_content = []
                            current_section_start_page = page_idx
                        else:
                            current_section_content.append(block_text)

                else:
                    # ---- Path B: scanned/image page — use OCR ----
                    logger.info(f"Page {page_idx} has no native text — running OCR.")
                    ocr_text = _ocr_page(page)
                    ocr_text = _strip_watermark(ocr_text)
                    if ocr_text:
                        page_text_parts.append(ocr_text)
                        current_section_content.append(ocr_text)
                        page_text_blocks.append(TextBlock(
                            text=ocr_text,
                            font_size=median_font_size,
                            is_bold=False,
                            bbox=(0, 0, page.rect.width, page.rect.height),
                            block_type="text"
                        ))

                p_text = "\n".join(page_text_parts)
                parsed_pages.append(ParsedPage(
                    page_number=page_idx,
                    text=TextNormalizer.normalize(p_text),
                    text_blocks=page_text_blocks,
                    has_tables=len(tables_in_page) > 0
                ))
                full_text_parts.append(p_text)

            # Append the last section
            if current_section_content or current_section_title:
                parsed_sections.append(ParsedSection(
                    section_id=f"S{section_id_counter}",
                    heading_text=current_section_title,
                    heading_level=current_section_level,
                    content=TextNormalizer.normalize("\n".join(current_section_content)),
                    start_page=current_section_start_page
                ))

            full_text = "\n\n".join(full_text_parts)
            normalized_full_text = TextNormalizer.normalize(full_text)

            # If still empty after OCR, raise
            if len(normalized_full_text.strip()) < 50:
                raise EmptyDocumentError("PDF document has no readable text content even after OCR.")

            # Try to grab title from metadata
            meta_title = doc.metadata.get("title") if doc.metadata else None

            return ParsedDocument(
                doc_id=doc_id,
                file_path=file_path,
                file_type="pdf",
                full_text=normalized_full_text,
                pages=parsed_pages,
                sections=parsed_sections,
                tables=parsed_tables,
                document_title=meta_title,
                page_count=len(doc),
                char_count=len(normalized_full_text)
            )
        finally:
            doc.close()


class DOCXParser:
    def parse(self, file_path: str, doc_id: str) -> ParsedDocument:
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            raise CorruptedDocumentError(f"Failed to open DOCX file: {e}")

        parsed_sections = []
        parsed_tables = []
        full_text_parts = []
        table_id_counter = 1
        section_id_counter = 1

        current_section_title = None
        current_section_level = 0
        current_section_content = []

        # In DOCX, we iterate logical elements of the body in order to preserve alignment
        body_elements = []
        
        # Traverse docx elements in order
        for child in doc.element.body:
            if child.tag.endswith('p'):
                body_elements.append(Paragraph(child, doc))
            elif child.tag.endswith('tbl'):
                body_elements.append(Table(child, doc))

        for element in body_elements:
            if isinstance(element, Paragraph):
                text = element.text
                if not text.strip():
                    continue

                style_name = element.style.name if element.style else "Normal"
                is_heading = style_name.startswith("Heading")
                heading_level = 0

                if is_heading:
                    match = re.search(r'\d+', style_name)
                    heading_level = int(match.group()) if match else 3

                full_text_parts.append(text)

                if is_heading:
                    # Save current section
                    if current_section_content or current_section_title:
                        parsed_sections.append(ParsedSection(
                            section_id=f"S{section_id_counter}",
                            heading_text=current_section_title,
                            heading_level=current_section_level,
                            content=TextNormalizer.normalize("\n".join(current_section_content)),
                            start_page=None
                        ))
                        section_id_counter += 1

                    current_section_title = text.strip()
                    current_section_level = heading_level
                    current_section_content = []
                else:
                    current_section_content.append(text)

            elif isinstance(element, Table):
                # Extract Table
                tbl_rows = []
                for row in element.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    tbl_rows.append(row_data)

                if not tbl_rows:
                    continue

                headers = tbl_rows[0]
                rows = tbl_rows[1:]
                raw_text = "\n".join([" | ".join(r) for r in tbl_rows])
                full_text_parts.append(raw_text)

                p_table = ParsedTable(
                    table_id=f"T{table_id_counter}",
                    headers=headers,
                    rows=rows,
                    source_page=None,
                    raw_text=raw_text
                )
                parsed_tables.append(p_table)
                table_id_counter += 1

                # Append table content to current section text
                current_section_content.append(raw_text)

        # Append last section
        if current_section_content or current_section_title:
            parsed_sections.append(ParsedSection(
                section_id=f"S{section_id_counter}",
                heading_text=current_section_title,
                heading_level=current_section_level,
                content=TextNormalizer.normalize("\n".join(current_section_content)),
                start_page=None
            ))

        full_text = "\n\n".join(full_text_parts)
        normalized_full_text = TextNormalizer.normalize(full_text)

        if not normalized_full_text.strip():
            raise EmptyDocumentError("DOCX document has no readable text content.")

        # Try to grab document title from core properties
        meta_title = None
        try:
            if doc.core_properties and doc.core_properties.title:
                meta_title = doc.core_properties.title
        except Exception:
            pass

        return ParsedDocument(
            doc_id=doc_id,
            file_path=file_path,
            file_type="docx",
            full_text=normalized_full_text,
            pages=[],  # DOCX doesn't have native pages
            sections=parsed_sections,
            tables=parsed_tables,
            document_title=meta_title,
            page_count=0,
            char_count=len(normalized_full_text)
        )


class DocumentParserModule:
    def __init__(self):
        self.pdf_parser = PDFParser()
        self.docx_parser = DOCXParser()

    def parse_document(self, file_path: str, doc_id: str) -> ParsedDocument:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        _, ext = os.path.splitext(file_path.lower())
        if ext == '.pdf':
            return self.pdf_parser.parse(file_path, doc_id)
        elif ext == '.docx':
            return self.docx_parser.parse(file_path, doc_id)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
