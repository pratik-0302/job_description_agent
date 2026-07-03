import unittest
import os
import shutil
import tempfile
import queue
import time
import datetime
import fitz
# pyrefly: ignore [missing-import]
import docx

from src.storage.sqlite_manager import SQLiteManager, DocumentRepository, Document
from src.ingestion.models import DocumentRecord, PipelineEvent, ParsedDocument
from src.ingestion.parser import DocumentParserModule, PDFParser, DOCXParser, TextNormalizer, DocumentProtectedError, ScannedDocumentError
from src.ingestion.discovery import DocumentDiscoveryModule, compute_sha256


class TestIngestionLayer(unittest.TestCase):
    def setUp(self):
        # Create a temp directory for the test files
        self.test_dir = tempfile.mkdtemp()
        self.db_path = os.path.join(self.test_dir, "test_metadata.db")
        
        # Setup SQLite
        self.db_manager = SQLiteManager(self.db_path)
        self.session = self.db_manager.get_session()
        self.repo = DocumentRepository(self.session)
        
        # Setup event queue and discovery module
        self.event_queue = queue.Queue()
        self.discovery = DocumentDiscoveryModule(
            watched_dir=self.test_dir,
            db_manager=self.db_manager,
            event_queue=self.event_queue
        )
        
        # Create mock PDF
        self.pdf_path = os.path.join(self.test_dir, "test_job.pdf")
        self._create_sample_pdf(self.pdf_path, "Software Engineer Intern", "We require Python and SQL.")
        
        # Create mock DOCX
        self.docx_path = os.path.join(self.test_dir, "test_job.docx")
        self._create_sample_docx(self.docx_path, "Data Scientist Role", "Required skills: Machine Learning, Python, and Statistics.")

    def tearDown(self):
        self.session.close()
        # Clean up database tables
        from src.storage.sqlite_manager import Base
        Base.metadata.drop_all(self.db_manager.engine)
        self.db_manager.engine.dispose()
        
        # Remove directory
        shutil.rmtree(self.test_dir)

    def _create_sample_pdf(self, path: str, title: str, body: str):
        doc = fitz.open()
        page = doc.new_page()
        # Insert Heading (large size)
        page.insert_text((50, 50), title, fontsize=20)
        # Insert normal text block
        page.insert_text((50, 100), body, fontsize=10)
        doc.save(path)
        doc.close()

    def _create_sample_docx(self, path: str, title: str, body: str):
        doc = docx.Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)
        doc.save(path)

    def test_database_crud(self):
        # Insert
        doc_record = {
            'doc_id': 'test-uuid-1',
            'file_path': '/path/to/file.pdf',
            'file_name': 'file.pdf',
            'file_extension': '.pdf',
            'file_size_bytes': 1024,
            'content_hash': 'abcde12345',
            'status': 'pending'
        }
        doc = self.repo.insert_document(doc_record)
        self.assertIsNotNone(doc)
        self.assertEqual(doc.status, 'pending')

        # Lookup
        found = self.repo.get_document_by_path('/path/to/file.pdf')
        self.assertIsNotNone(found)
        self.assertEqual(found.content_hash, 'abcde12345')

        # Update
        self.repo.update_document_status('test-uuid-1', 'indexed', failure_reason=None)
        found = self.repo.get_document_by_path('/path/to/file.pdf')
        self.assertEqual(found.status, 'indexed')
        self.assertIsNotNone(found.indexed_at)

        # List
        all_docs = self.repo.list_documents()
        self.assertEqual(len(all_docs), 1)

        # Delete
        self.repo.delete_document('test-uuid-1')
        found = self.repo.get_document_by_path('/path/to/file.pdf')
        self.assertIsNone(found)

    def test_file_hashing_and_scans(self):
        # Run force_rescan on startup scan
        self.discovery.force_rescan()
        
        # Verify two items are registered
        docs = self.repo.list_documents()
        self.assertEqual(len(docs), 2)
        
        # Verify events are queued
        events = []
        while not self.event_queue.empty():
            events.append(self.event_queue.get())
        
        self.assertEqual(len(events), 2)
        self.assertTrue(any(e.event_type == "new" and e.file_path == self.pdf_path for e in events))
        self.assertTrue(any(e.event_type == "new" and e.file_path == self.docx_path for e in events))

    def test_duplicate_detection(self):
        # Create a copy of the PDF (same content, different filename)
        duplicate_path = os.path.join(self.test_dir, "test_job_copy.pdf")
        shutil.copyfile(self.pdf_path, duplicate_path)

        # Scan
        self.discovery.force_rescan()

        # The document repo should still only have 2 unique entries because the 3rd is a content hash duplicate
        docs = self.repo.list_documents()
        self.assertEqual(len(docs), 2)

    def test_pdf_docx_parsers(self):
        parser_module = DocumentParserModule()

        # Parse PDF
        parsed_pdf = parser_module.parse_document(self.pdf_path, "doc-pdf-1")
        self.assertEqual(parsed_pdf.file_type, "pdf")
        self.assertIn("Software Engineer Intern", parsed_pdf.full_text)
        self.assertIn("Python", parsed_pdf.full_text)
        self.assertEqual(len(parsed_pdf.pages), 1)
        self.assertEqual(parsed_pdf.pages[0].page_number, 1)

        # Parse DOCX
        parsed_docx = parser_module.parse_document(self.docx_path, "doc-docx-1")
        self.assertEqual(parsed_docx.file_type, "docx")
        self.assertIn("Data Scientist Role", parsed_docx.full_text)
        self.assertIn("Machine Learning", parsed_docx.full_text)
        self.assertEqual(len(parsed_docx.sections), 1)
        self.assertEqual(parsed_docx.sections[0].heading_text, "Data Scientist Role")

    def test_text_normalization(self):
        raw_text = "Hello \t world.\r\n•  Required Skill:\n- Python\n* SQL\n\n\n\nNew section."
        normalized = TextNormalizer.normalize(raw_text)
        self.assertIn("Hello world.", normalized)
        self.assertIn("• Required Skill:", normalized)
        self.assertIn("• Python", normalized)
        self.assertIn("• SQL", normalized)
        # Ensure debounced newlines
        self.assertNotIn("\n\n\n\n", normalized)

    def test_watchdog_event_handling(self):
        # Reset queue
        while not self.event_queue.empty():
            self.event_queue.get()
            
        # Simulate creating a file
        new_file = os.path.join(self.test_dir, "new_temp.pdf")
        self._create_sample_pdf(new_file, "New Software Role", "Requires Go.")
        
        self.discovery.handle_file_created(new_file)
        
        # Verify event queued
        self.assertFalse(self.event_queue.empty())
        evt = self.event_queue.get()
        self.assertEqual(evt.event_type, "new")
        self.assertEqual(evt.file_path, new_file)


if __name__ == "__main__":
    unittest.main()
